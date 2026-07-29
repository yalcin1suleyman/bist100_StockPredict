"""
Doğrulama Scripti: Kullanıcının Drive'dan indirdiği "bist_100 (1).docx" dosyasını kontrol eder.
"""
import docx

FILE = r"c:\stajProje\dosyalar\bist_100 (1).docx"
doc = docx.Document(FILE)
full_text = "\n".join([p.text for p in doc.paragraphs])

print("=" * 60)
print("USER FILE KONTROL RAPORU")
print("=" * 60)

checks = [
    ("1a: Bolum 3.4 veri bolme %70/%15/%15",
     "%70 eğitim (train), %15 doğrulama (validation) ve %15 test (test) seti olacak şekilde",
     True),
    
    ("1b: DL validation cumlesi guncellendi",
     "derin öğrenme modellerinde doğrulama seti üzerinde",
     True),
    
    ("1c: Bolum 4 %70/%15/%15",
     "%70 eğitim, %15 doğrulama (validation) ve %15 test kümelerine ayrılmıştır",
     True),
    
    ("1d: Bolum 5.1.1 %70/%15/%15",
     "%70'lik dilimi eğitim, %15'lik dilimi doğrulama ve son %15'lik dilimi ise test seti",
     True),
    
    ("1e: Bolum 5.1 %70/%15/%15",
     "%70 eğitim, %15 doğrulama ve %15 test kümeleri kronolojik",
     True),
    
    ("2: Bolum 3.4 kisaltildi (Eksik Veri Paragrafi)",
     "Bölüm 5.1.3'te sunulmuştur",
     True),
    
    ("3: Multicollinearity eklendi",
     "çoklu doğrusal bağlantı (multicollinearity) sorunudur",
     True),
    
    ("4: Sekil 5.13'te",
     "Şekil 5.13'te XGBoost",
     True),
    
    ("5: Kesilmis cumle tamamlandi",
     "CNN-LSTM ve Transformer modellerinin öğrenme eğrileri incelenmiştir",
     True),
    
    ("6: SHAP yazim hatasi duzeltildi",
     "SHAP analizleri yalnızca matematiksel",
     True),
    
    ("7: 5.10 Model Karmasikligi",
     "5.10 Model Karmaşıklığı",
     True),
]

# Olmamasi gereken seyler
negative_checks = [
    ("Eski %80/%20 eğitim/test ifadesi kalmasin",
     "%80 eğitim",
     False),
    
    ("Eski %10 dogrulama kalmasin",
     "%10 doğrulama (validation) ve %20 test",
     False),
    
    ("Eski %80'lik dilimi ifadesi kalmasin",
     "ilk %80'lik dilimi eğitim, son %20'lik dilimi ise test",
     False),
    
    ("Eski uzun Eksik Veri cumleleri kalmasin",
     "Tablo 3.1'de görüldüğü üzere, özellikle VIX endeksinde uluslararası tatillerin",
     False),
    
    ("Kesilmis CNN-LST kalmasin (em-dash versiyonu)",
     "CNN\u2013LST",
     False),
    
    ("5.12 numaralandirma kalmasin",
     "5.12 Model Karmaşıklığı",
     False),
]

all_pass = True

for desc, text, should_exist in checks:
    found = text in full_text
    # 'CNN-LSTM' check is special since text might have em-dash "CNN–LSTM" based on user's manual edit
    if not found and "CNN" in text and "LSTM" in text:
        found = "CNN–LSTM ve Transformer modellerinin öğrenme eğrileri incelenmiştir" in full_text or \
                "CNN-LSTM ve Transformer modellerinin öğrenme eğrileri incelenmiştir" in full_text
                
    if found == should_exist:
        print(f"  [OK] {desc}")
    else:
        all_pass = False
        print(f"  [FAIL] {desc} - BEKLENEN DURUM SAGLANAMADI")

print()
for desc, text, should_exist in negative_checks:
    found = text in full_text
    
    # Exclude matches related to "LightGBM (%80 eğitim...)" in section 5.1.3 since we didn't remove that
    if found and text == "%80 eğitim":
        found = False
        for p in doc.paragraphs:
            if "%80 eğitim" in p.text and "LightGBM" not in p.text:
                found = True
                break

    if found == should_exist:
        print(f"  [OK] {desc}")
    else:
        all_pass = False
        print(f"  [FAIL] {desc} - Eski metin hala belgede!")

print("\n" + "=" * 60)
if all_pass:
    print("TUM KONTROLLER BASARILI! Kullanicinin dosyasi tamamen dogru.")
else:
    print("BAZI KONTROLLER BASARISIZ! Lutfen FAIL satirlarini inceleyin.")
print("=" * 60)
